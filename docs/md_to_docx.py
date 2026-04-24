from pathlib import Path

from docx import Document


def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    lines = md_path.read_text(encoding="utf-8").splitlines()

    in_code_block = False
    for raw in lines:
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line if line else " ")
            run.font.name = "Consolas"
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue

        stripped = line.lstrip()
        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue
        if stripped.startswith(("- [ ] ", "- [x] ", "- [X] ")):
            text = stripped[6:].strip()
            doc.add_paragraph(text, style="List Bullet")
            continue

        doc.add_paragraph(line)

    doc.save(str(docx_path))


if __name__ == "__main__":
    root = Path(__file__).resolve().parent
    md_file = root / "HANDOVER.md"
    docx_file = root / "HANDOVER.docx"
    convert_md_to_docx(md_file, docx_file)
    print(f"Created: {docx_file}")
